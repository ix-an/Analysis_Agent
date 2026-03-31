from langchain.tools import tool
from pydantic import BaseModel, Field
from docx import Document
import os
import pandas as pd


class FilePathArgs(BaseModel):
    path: str = Field(..., description="文件的绝对路径")


@tool(args_schema=FilePathArgs)
def file_reader_tool(path: str) -> str:
    """
    支持读取 Word(.docx), Excel(.xlsx/.xls), CSV(.csv) 以及文本文件内容。
    读取时会自动删除重复行，并将缺失值填充为 'None'。
    """
    ext = os.path.splitext(path)[-1].lower()
    content = ""

    try:
        # 1. 处理 Excel/CSV (优先使用 pandas 进行自动化清洗)
        if ext in [".xlsx", ".xls", ".csv"]:
            df = pd.read_csv(path) if ext == ".csv" else pd.read_excel(path)

            # --- 核心逻辑：自动清洗 ---
            df = df.drop_duplicates()  # 自动删除重复值
            df = df.fillna("None")  # 自动填充缺失值为 None

            # 转换为 Markdown 格式，方便 LLM 理解数据结构
            content = f"【表格数据概览（已完成去重和缺失值填充）】\n{df.to_markdown()}"

        # 2. 处理 Word
        elif ext == ".docx":
            doc = Document(path)
            # 读取段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # 读取表格
            tables_data = []
            for table in doc.tables:
                for row in table.rows:
                    tables_data.append(
                        " | ".join([cell.text.strip() if cell.text.strip() else "None" for cell in row.cells]))

            # 去重处理
            unique_paragraphs = list(dict.fromkeys(paragraphs))
            content = "【文档正文】\n" + "\n".join(unique_paragraphs) + "\n\n【文档表格】\n" + "\n".join(tables_data)

        # 3. 处理普通文本
        else:
            with open(path, 'r', encoding='utf-8') as f:
                content = f.read()

        return f"文件内容读取并清洗成功：\n\n{content}"

    except Exception as e:
        return f"读取文件出错: {str(e)}"