import os
from langchain.tools import tool
from pydantic import BaseModel, Field
from app.utils.Logger import Logger
from docx import Document
import time
from pathlib import Path

# 定义参数
class Args(BaseModel):
    content: str = Field(..., description="分析报告内容")
logger = Logger.get_logger(__name__)

@tool("docx_write_tool", args_schema=Args)
def docx_write_tool(content:str)-> str:
    """
    生成分析报告docx文件，并返回可下载链接
    """
    # 1. 创建文档
    doc = Document()
    doc.add_heading('文件分析报告', level=1)
    doc.add_paragraph(content)

    # 2. ✅ 核心路径：指向 app/static/download
    file_name = time.strftime("%Y%m%d%H%M%S", time.localtime())
    base_dir = Path(__file__).resolve().parent.parent.parent  # 项目根目录
    save_dir = base_dir / "app" / "static" / "download"  # 强制加app层级

    # 3. 自动创建文件夹（不存在则创建）
    save_dir.mkdir(parents=True, exist_ok=True)

    # 4. 保存文件
    file_path = save_dir / f"{file_name}.docx"
    doc.save(file_path)

    # 5. 返回下载链接
    down_url = f"http://localhost:8000/static/download/{file_name}.docx"
    logger.info(f"报告已生成：{file_path}")
    logger.info(f"下载链接：{down_url}")
    return f"完整报告请下载：{down_url}"